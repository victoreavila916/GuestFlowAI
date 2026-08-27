from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUT = "/Users/victoravila/Desktop/guestflow-ai/Property_Operations_Information_Template.docx"

GREEN = "1E4B3A"
MOSS = "377057"
INK = "18231F"
MUTED = "5C6961"
PALE = "EEF5EE"
LINE = "C9D8CB"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=LINE, size="8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=100, start=150, bottom=100, end=150):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_width(cell, dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(dxa))
    tc_w.set(qn("w:type"), "dxa")


def keep_row_together(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def add_run(paragraph, text, size=10.5, color=INK, bold=False, italic=False):
    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic
    return run


def set_para(paragraph, before=0, after=6, line=1.25):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_label(doc, label, helper, lines=2):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    keep_row_together(table.rows[0])
    cell = table.cell(0, 0)
    set_width(cell, 9360)
    set_cell_border(cell)
    set_cell_margins(cell)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    p = cell.paragraphs[0]
    set_para(p, after=2, line=1.1)
    add_run(p, label, size=10.5, color=GREEN, bold=True)
    p = cell.add_paragraph()
    set_para(p, after=7, line=1.15)
    add_run(p, helper, size=8.5, color=MUTED, italic=True)
    for _ in range(lines):
        p = cell.add_paragraph()
        set_para(p, after=5, line=1.15)
        add_run(p, "________________________________________________________________________________", size=9, color="B6C3B7")
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_section(doc, number, title, intro):
    p = doc.add_paragraph()
    set_para(p, before=14, after=3, line=1.0)
    p.paragraph_format.keep_with_next = True
    add_run(p, f"{number:02d}", size=10, color=MOSS, bold=True)
    add_run(p, f"  {title}", size=15, color=GREEN, bold=True)
    p = doc.add_paragraph()
    set_para(p, after=8, line=1.15)
    p.paragraph_format.keep_with_next = True
    add_run(p, intro, size=9.5, color=MUTED)


def add_header_footer(section):
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para(p, after=0, line=1.0)
    add_run(p, "GUESTFLOW  |  PROPERTY OPERATIONS INFORMATION", size=8.5, color=MOSS, bold=True)
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(p, before=0, after=0, line=1.0)
    add_run(p, "Complete this form in plain language. If a detail does not apply, write “N/A.”", size=8, color=MUTED)


def add_callout(doc):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.cell(0, 0)
    set_width(cell, 9360)
    set_cell_shading(cell, PALE)
    set_cell_border(cell, color="B9D3BC")
    set_cell_margins(cell, top=130, start=170, bottom=130, end=170)
    p = cell.paragraphs[0]
    set_para(p, after=3, line=1.15)
    add_run(p, "What not to repeat", size=10, color=GREEN, bold=True)
    p = cell.add_paragraph()
    set_para(p, after=0, line=1.2)
    add_run(p, "Do not include the property name, street address, Wi-Fi name/password, check-in or check-out time, house rules, or emergency contact. GuestFlow already collects those fields separately.", size=9.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    add_header_footer(section)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)

    p = doc.add_paragraph()
    set_para(p, before=10, after=3, line=1.0)
    add_run(p, "GUESTFLOW", size=10, color=MOSS, bold=True)
    p = doc.add_paragraph()
    set_para(p, after=4, line=1.0)
    add_run(p, "Property Operations Information", size=26, color=GREEN, bold=True)
    p = doc.add_paragraph()
    set_para(p, after=13, line=1.25)
    add_run(p, "A refined owner-completion form for the details guests ask about most. These answers help GuestFlow provide accurate, useful guidance without making things up.", size=11, color=MUTED)
    add_callout(doc)

    p = doc.add_paragraph()
    set_para(p, before=4, after=7, line=1.15)
    add_run(p, "Completion standard", size=10, color=GREEN, bold=True)
    add_run(p, "  •  Be specific enough that a first-time guest can act without contacting you. Include locations, simple steps, and exceptions. Avoid “see manual” or “use as normal.”", size=9.5, color=INK)

    add_section(doc, 1, "Arrival, Access & Parking", "Focus on the practical steps between arriving nearby and being comfortably inside.")
    add_label(doc, "Building, gate, elevator, or shared-entry instructions", "Include codes, call-box names, access app steps, and any timing restrictions.", 3)
    add_label(doc, "Parking plan", "Where to park; space/permit details; garage opener location; street-parking limits; towing risks; EV charging instructions.", 3)
    add_label(doc, "Keys, locks, and access backup", "Where physical keys, fobs, remotes, or lockboxes are located; what to do if access fails.", 3)
    add_label(doc, "Arrival accessibility notes", "Stairs, uneven paths, lighting, door widths, elevators, or anything a guest should plan for before arrival.", 2)

    add_section(doc, 2, "Home Systems & Essential Equipment", "Tell guests what is unusual, location-specific, or easy to operate incorrectly.")
    add_label(doc, "Heating, air conditioning & thermostat", "Thermostat location; recommended settings; whether windows/doors must stay closed; any systems guests must not adjust.", 3)
    add_label(doc, "Hot water, plumbing & shutoffs", "How long hot water takes, any faucet/toilet quirks, and the location of the main water shutoff for urgent leaks.", 3)
    add_label(doc, "Electrical panel, appliances & special equipment", "Breaker location and labels; appliance instructions; fireplace, generator, sauna, hot tub, pool, grill, or other equipment. State what is off-limits.", 4)
    add_label(doc, "Laundry", "Washer/dryer location, detergent provided, operating steps, lint filter, and any restrictions.", 2)

    add_section(doc, 3, "Safety, Health & Property Protection", "Capture location-based safety details that help guests act quickly and avoid damage.")
    add_label(doc, "Safety equipment locations", "Smoke/CO alarms, fire extinguishers, first-aid kit, flashlights, fire blanket, and emergency exits. Include exact locations.", 3)
    add_label(doc, "Hazards or seasonal risks", "Stairs, slippery surfaces, wildlife, water features, severe-weather procedures, snow/ice, beach or mountain risks, or construction nearby.", 3)
    add_label(doc, "Damage prevention", "Anything guests should know to prevent damage: septic limits, disposal limits, delicate floors, window treatments, plumbing sensitivities, or furniture care.", 3)

    add_section(doc, 4, "Daily Living, Cleaning & Waste", "Give the concise operational information guests need while staying in the home.")
    add_label(doc, "Trash, recycling & bulk waste", "Bin location, pickup days, sorting rules, bag requirements, and what to do with extra trash.", 3)
    add_label(doc, "Kitchen essentials & appliances", "Coffee maker type, filters/pods, ice maker, dishwasher basics, water filter, and the location of common cooking essentials.", 3)
    add_label(doc, "Entertainment & workspace", "TV remotes, streaming sign-in guidance, speaker instructions, desk/workspace location, printer, or network equipment guests should not unplug.", 3)
    add_label(doc, "Supplies, cleaning & pest response", "Where to find extra linens, towels, paper goods, cleaning supplies, and what guests should do if they see pests or need a refill.", 3)

    add_section(doc, 5, "Neighborhood & Local Orientation", "Include only recommendations and restrictions that materially improve a guest’s stay.")
    add_label(doc, "Nearest essentials", "Closest grocery, pharmacy, gas/EV charging, ATM, urgent care, and late-night option. Include names and approximate travel times.", 3)
    add_label(doc, "Transportation & local navigation", "Transit stops, rideshare pickup spot, airport/rail guidance, bike storage, local driving hazards, and parking restrictions near key destinations.", 3)
    add_label(doc, "Curated local recommendations", "A short list of 3–5 reliable restaurants, coffee shops, family activities, or experiences. Note reservations, seasonality, or distance where useful.", 4)

    add_section(doc, 6, "Host Notes for GuestFlow", "Use this section for the questions guests ask repeatedly or the details that are easy to miss.")
    add_label(doc, "Frequently asked questions", "List the question and the exact answer you want GuestFlow to give. Include any local wording or preferred guidance.", 5)
    add_label(doc, "Escalation and service notes", "When should GuestFlow tell a guest to contact the host rather than attempting an answer? Include maintenance boundaries, vendor access, and anything that needs human approval.", 4)
    add_label(doc, "Anything else a thoughtful guest should know", "Include the small, high-value details that make the stay feel easy and well-hosted.", 4)

    doc.save(OUT)


if __name__ == "__main__":
    main()
